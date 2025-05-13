import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Mm
import io
from datetime import datetime

class ImageToA4App:
    def __init__(self, root):
        self.root = root
        self.root.title("Фото на A4 с поворотом")
        self.root.geometry("1100x700")
        self.root.minsize(1000, 600)

        # Переменные
        self.images = []  # Пути к оригинальным изображениям
        self.rotations = []  # Углы поворота для каждого изображения
        self.previews = []
        self.margin_mm = 10
        self.a4_width_mm = 210
        self.a4_height_mm = 297

        # Создаем основной интерфейс
        self.create_widgets()

    def create_widgets(self):
        # Фрейм для управления
        control_frame = tk.Frame(self.root)
        control_frame.pack(pady=10, fill=tk.X)

        # Кнопки управления
        btn_select = tk.Button(control_frame, text="Выбрать изображения", command=self.select_images)
        btn_select.pack(side=tk.LEFT, padx=5)

        btn_up = tk.Button(control_frame, text="Вверх", command=self.move_up)
        btn_up.pack(side=tk.LEFT, padx=5)

        btn_down = tk.Button(control_frame, text="Вниз", command=self.move_down)
        btn_down.pack(side=tk.LEFT, padx=5)

        btn_remove = tk.Button(control_frame, text="Удалить", command=self.remove_selected)
        btn_remove.pack(side=tk.LEFT, padx=5)

        btn_rotate = tk.Button(control_frame, text="Повернуть", command=self.rotate_image)
        btn_rotate.pack(side=tk.LEFT, padx=5)

        btn_process = tk.Button(control_frame, text="Создать DOCX", command=self.process_images)
        btn_process.pack(side=tk.LEFT, padx=5)

        # Основной контейнер с двумя панелями
        main_container = tk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        main_container.pack(fill=tk.BOTH, expand=True)

        # Левая панель - список изображений
        left_panel = tk.Frame(main_container)
        main_container.add(left_panel)

        scrollbar = tk.Scrollbar(left_panel)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.image_list = tk.Listbox(
            left_panel,
            yscrollcommand=scrollbar.set,
            selectmode=tk.SINGLE,
            width=40,
            height=30
        )
        self.image_list.pack(fill=tk.BOTH, expand=True)
        self.image_list.bind('<<ListboxSelect>>', self.update_preview)

        scrollbar.config(command=self.image_list.yview)

        # Правая панель - предпросмотр
        right_panel = tk.Frame(main_container)
        main_container.add(right_panel)

        self.preview_label = tk.Label(right_panel, text="Предпросмотр будет здесь", bg='white')
        self.preview_label.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Статус бар
        self.status_bar = tk.Label(self.root, text="Готово", bd=1, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(fill=tk.X)

    def select_images(self):
        file_paths = filedialog.askopenfilenames(
            title="Выберите изображения",
            filetypes=[("Image files", "*.jpg *.jpeg *.png *.bmp")]
        )
        if file_paths:
            self.image_list.delete(0, tk.END)
            self.images = []
            self.rotations = []
            for path in file_paths:
                self.image_list.insert(tk.END, path.split('/')[-1])
                self.images.append(path)
                self.rotations.append(0)  # Начальный угол поворота
            self.update_preview()
            self.update_status(f"Загружено {len(file_paths)} изображений")

    def move_up(self):
        selected = self.image_list.curselection()
        if not selected or selected[0] == 0:
            return
        pos = selected[0]
        # Перемещаем в списках
        self.images[pos], self.images[pos - 1] = self.images[pos - 1], self.images[pos]
        self.rotations[pos], self.rotations[pos - 1] = self.rotations[pos - 1], self.rotations[pos]
        # Перемещаем в Listbox
        item = self.image_list.get(pos)
        self.image_list.delete(pos)
        self.image_list.insert(pos - 1, item)
        self.image_list.select_set(pos - 1)
        self.update_preview()

    def move_down(self):
        selected = self.image_list.curselection()
        if not selected or selected[0] == len(self.images) - 1:
            return
        pos = selected[0]
        # Перемещаем в списках
        self.images[pos], self.images[pos + 1] = self.images[pos + 1], self.images[pos]
        self.rotations[pos], self.rotations[pos + 1] = self.rotations[pos + 1], self.rotations[pos]
        # Перемещаем в Listbox
        item = self.image_list.get(pos)
        self.image_list.delete(pos)
        self.image_list.insert(pos + 1, item)
        self.image_list.select_set(pos + 1)
        self.update_preview()

    def remove_selected(self):
        selected = self.image_list.curselection()
        if not selected:
            return
        pos = selected[0]
        self.image_list.delete(pos)
        del self.images[pos]
        del self.rotations[pos]
        self.update_preview()
        self.update_status(f"Удалено изображение. Осталось: {len(self.images)}")

    def rotate_image(self):
        selected = self.image_list.curselection()
        if not selected:
            return
        pos = selected[0]
        # Увеличиваем угол поворота на 90 градусов
        self.rotations[pos] = (self.rotations[pos] + 90) % 360
        self.update_preview()
        self.update_status(f"Изображение повернуто на {self.rotations[pos]}°")

    def update_preview(self, event=None):
        if not self.images:
            self.preview_label.config(image='', text="Нет изображений для предпросмотра")
            return

        selected = self.image_list.curselection()
        if not selected:
            return

        pos = selected[0]
        img_path = self.images[pos]
        rotation = self.rotations[pos]

        try:
            # Загружаем изображение
            img = Image.open(img_path)

            # Применяем поворот
            if rotation != 0:
                img = img.rotate(-rotation, expand=True)  # Отрицательный угол для поворота по часовой

            # Рассчитываем размеры для предпросмотра
            preview_width = 400
            margin_px = int(self.margin_mm * preview_width / self.a4_width_mm)
            img_width = preview_width - 2 * margin_px
            aspect_ratio = img.height / img.width
            img_height = int(img_width * aspect_ratio)

            # Создаем изображение A4 для предпросмотра
            a4_preview = Image.new('RGB', (preview_width, int(preview_width * self.a4_height_mm / self.a4_width_mm)),
                                   'white')

            # Масштабируем изображение
            img_resized = img.resize((img_width, img_height), Image.LANCZOS)

            # Размещаем изображение на A4
            x = margin_px
            y = (a4_preview.height - img_resized.height) // 2
            a4_preview.paste(img_resized, (x, y))

            # Конвертируем для Tkinter
            photo = ImageTk.PhotoImage(a4_preview)

            # Обновляем label
            self.preview_label.config(image=photo)
            self.preview_label.image = photo
            self.preview_label.text = ''

            # Обновляем статус
            self.update_status(
                f"Предпросмотр: {img_path.split('/')[-1]} ({pos + 1}/{len(self.images)}), Поворот: {rotation}°")

        except Exception as e:
            self.preview_label.config(image='', text=f"Ошибка загрузки изображения: {str(e)}")
            self.update_status(f"Ошибка: {str(e)}")

    def process_images(self):
        if not self.images:
            messagebox.showerror("Ошибка", "Нет изображений для обработки")
            return


        output_file = filedialog.asksaveasfilename(
            title="Сохранить как DOCX",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            initialfile="file_" + datetime.now().strftime("%d%m%Y%H%M%S")
        )
        if not output_file:
            return

        doc = Document()

        for i, (img_path, rotation) in enumerate(zip(self.images, self.rotations)):
            try:
                img = Image.open(img_path)

                # Применяем поворот перед вставкой в документ
                if rotation != 0:
                    img = img.rotate(-rotation, expand=True)

                width_mm = self.a4_width_mm - 2 * self.margin_mm
                aspect_ratio = img.height / img.width
                height_mm = width_mm * aspect_ratio

                if i > 0:
                    doc.add_section()

                section = doc.sections[i]
                section.page_height = Mm(self.a4_height_mm)
                section.page_width = Mm(self.a4_width_mm)
                section.top_margin = Mm(self.margin_mm)
                section.bottom_margin = Mm(self.margin_mm)
                section.left_margin = Mm(self.margin_mm)
                section.right_margin = Mm(self.margin_mm)

                # Сохраняем временное изображение с поворотом
                temp_img = io.BytesIO()
                img.save(temp_img, format='PNG')
                temp_img.seek(0)

                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(temp_img, width=Mm(width_mm))

                if i < len(self.images) - 1:
                    doc.add_page_break()

            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось обработать {img_path}: {str(e)}")
                continue

        try:
            doc.save(output_file)
            messagebox.showinfo("Готово", f"Файл успешно сохранен:\n{output_file}")
            self.update_status(f"DOCX создан: {output_file}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {str(e)}")
            self.update_status(f"Ошибка сохранения: {str(e)}")

    def update_status(self, message):
        self.status_bar.config(text=message)


if __name__ == "__main__":
    root = tk.Tk()
    app = ImageToA4App(root)
    root.mainloop()